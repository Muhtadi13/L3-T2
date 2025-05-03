
#####################################################################################
# import pandas as pd
# from docx import Document
# from tabulate import tabulate  # For better table formatting in console output

# # Load the data
# file_path = "TSP/tsp9990.csv"  # Replace with the actual file path
# data = pd.read_csv(file_path)

# # Extract methods
# constructive_methods = ['Nearest Neighbour', 'Nearest Insertion', 'Cheapest Insertion']
# perturbative_methods = ['2-OPT', '3-OPT', 'Node Shift']

# # Calculate total experiments (assuming all rows are part of the experiment)
# total_experiments = 21

# # Filter data for constructive and perturbative methods
# constructive_data = data[data['Construction Heuristics'].isin(constructive_methods)]
# perturbative_data = data[data['Improvement Heuristics'].isin(perturbative_methods)]

# # Calculate average costs for constructive methods
# constructive_avg_costs = (
#     constructive_data.groupby('Construction Heuristics')['Cost']
#     .mean()
#     .reset_index()
#     .rename(columns={'Cost': 'Average Cost'})
# )

# # Calculate average costs for perturbative methods
# perturbative_avg_costs = (
#     perturbative_data.groupby('Improvement Heuristics')['Cost']
#     .mean()
#     .reset_index()
#     .rename(columns={'Cost': 'Average Cost'})
# )

# # Save the results in a Word document
# doc = Document()
# doc.add_heading("TSP Analysis: Average Costs for Each Method", level=1)

# # Add constructive methods table
# doc.add_heading("Constructive Methods", level=2)
# doc.add_paragraph("The table below shows the average costs for constructive methods.")
# constructive_table = doc.add_table(rows=1, cols=2)
# constructive_table.style = 'Table Grid'

# # Add header
# constructive_table.cell(0, 0).text = "Construction Heuristic"
# constructive_table.cell(0, 1).text = "Average Cost"

# # Add data rows
# for _, row in constructive_avg_costs.iterrows():
#     cells = constructive_table.add_row().cells
#     cells[0].text = row['Construction Heuristics']
#     cells[1].text = f"{row['Average Cost']:.2f}"

# # Add perturbative methods table
# doc.add_heading("Perturbative Methods", level=2)
# doc.add_paragraph("The table below shows the average costs for perturbative methods.")
# perturbative_table = doc.add_table(rows=1, cols=2)
# perturbative_table.style = 'Table Grid'

# # Add header
# perturbative_table.cell(0, 0).text = "Improvement Heuristic"
# perturbative_table.cell(0, 1).text = "Average Cost"

# # Add data rows
# for _, row in perturbative_avg_costs.iterrows():
#     cells = perturbative_table.add_row().cells
#     cells[0].text = row['Improvement Heuristics']
#     cells[1].text = f"{row['Average Cost']:.2f}"

# # Save the document
# doc.save("TSP_Average_Costs.docx")

# print("Document created successfully: TSP_Average_Costs.docx")

# # Print results to console for reference
# print("\nConstructive Methods Average Costs:")
# print(tabulate(constructive_avg_costs, headers="keys", tablefmt="pretty", showindex=False))

# print("\nPerturbative Methods Average Costs:")
# print(tabulate(perturbative_avg_costs, headers="keys", tablefmt="pretty", showindex=False))

# Load the data
import pandas as pd
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
file_path = "TSP/tsp9990.csv"  # Replace with the actual file path
data = pd.read_csv(file_path)

# Extract the node number from the filename
def extract_node(filename):
    match = re.search(r'(\d+)\.tsp$', filename)
    return int(match.group(1)) if match else None

data['Node'] = data['Filename'].apply(extract_node)

# Handle cases where the node number could not be extracted
if data['Node'].isnull().any():
    raise ValueError("Some filenames in the data are improperly formatted or missing node numbers.")

# Identify constructive and perturbative methods
constructive_methods = ['Nearest Neighbour', 'Nearest Insertion', 'Cheapest Insertion']
perturbative_methods = ['2-OPT', '3-OPT', 'Node Shift']

# Calculate deviation for constructive methods
constructive_avg = (
    data[data['Improvement Heuristics'] == 'None']
    .groupby('Node')['Cost']
    .mean()
    .reset_index()
)
constructive_avg.rename(columns={'Cost': 'Avg Cost (Constructive)'}, inplace=True)

# Calculate deviation for perturbative methods
perturbative_avg = (
    data[data['Improvement Heuristics'].isin(perturbative_methods)]
    .groupby(['Node', 'Improvement Heuristics'])['Cost']
    .mean()
    .reset_index()
)

# Create a document
doc = Document()
doc.add_heading("TSP Analysis: Average Deviation per Node", level=1)

# Add constructive average deviation to the document
doc.add_heading("Constructive Methods", level=2)
for _, row in constructive_avg.iterrows():
    doc.add_paragraph(f"Node: {row['Node']}, Avg Cost: {row['Avg Cost (Constructive)']}")

# Add perturbative average deviation to the document
doc.add_heading("Perturbative Methods", level=2)
for method in perturbative_methods:
    doc.add_heading(f"Method: {method}", level=3)
    method_data = perturbative_avg[perturbative_avg['Improvement Heuristics'] == method]
    for _, row in method_data.iterrows():
        doc.add_paragraph(f"Node: {row['Node']}, Avg Cost: {row['Cost']}")

# Plot graphs
# Constructive methods
plt.figure(figsize=(10, 6))
plt.plot(constructive_avg['Node'], constructive_avg['Avg Cost (Constructive)'], marker='o', label='Constructive Avg Cost')
plt.title("Constructive Methods: Average Cost vs Node Number")
plt.xlabel("Node Number")
plt.ylabel("Average Cost")
plt.grid()
plt.legend()
constructive_plot_path = "constructive_avg_cost.png"
plt.savefig(constructive_plot_path)
plt.close()
doc.add_picture(constructive_plot_path, width=Inches(5))

# Perturbative methods
plt.figure(figsize=(10, 6))
for method in perturbative_methods:
    method_data = perturbative_avg[perturbative_avg['Improvement Heuristics'] == method]
    plt.plot(method_data['Node'], method_data['Cost'], marker='o', label=method)
plt.title("Perturbative Methods: Average Cost vs Node Number")
plt.xlabel("Node Number")
plt.ylabel("Average Cost")
plt.grid()
plt.legend()
perturbative_plot_path = "perturbative_avg_cost.png"
plt.savefig(perturbative_plot_path)
plt.close()
doc.add_picture(perturbative_plot_path, width=Inches(5))

# Save the document
doc.save("TSP_Average_Cost_Analysis.docx")

print("Document created successfully: TSP_Average_Cost_Analysis.docx")