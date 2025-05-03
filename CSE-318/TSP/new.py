import pandas as pd
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Load the data
file_path = "TSP/tsp9990.csv"  # Replace with the actual file path
data = pd.read_csv(file_path)

# Extract the node number from the filename
def extract_node(filename):
    match = re.search(r'(\d+)\.tsp$', filename)
    return int(match.group(1)) if match else None

data['Node'] = data['Filename'].apply(extract_node)

# Handle cases where the node number could not be extracted
if data['Node'].isnull().any():
    raise ValueError("Some filenames in the data are improperly formatted or missing node numbers.")

# Identify constructive and perturbative methods
constructive_methods = ['Nearest Neighbour', 'Nearest Insertion', 'Cheapest Insertion']
perturbative_methods = ['2-OPT', '3-OPT', 'Node Shift']

# Calculate average improvement time (Improvement Time - Construction Time)
data['Improvement Time (Difference)'] = data['Improvement Time(ms)'] - data['Construction Time(ms)']

# Calculate average improvement time for constructive methods
constructive_avg_time = (
    data[data['Construction Heuristics'].isin(constructive_methods)]
    .groupby(['Node', 'Construction Heuristics'])['Improvement Time (Difference)']
    .mean()
    .reset_index()
)

# Calculate average improvement time for perturbative methods
perturbative_avg_time = (
    data[data['Improvement Heuristics'].isin(perturbative_methods)]
    .groupby(['Node', 'Improvement Heuristics'])['Improvement Time (Difference)']
    .mean()
    .reset_index()
)

# Create a document
doc = Document()
doc.add_heading("TSP Analysis: Average Improvement Time per Node", level=1)

# Add constructive average improvement time to the document
doc.add_heading("Constructive Methods", level=2)
for method in constructive_methods:
    doc.add_heading(f"Method: {method}", level=3)
    method_data = constructive_avg_time[constructive_avg_time['Construction Heuristics'] == method]
    for _, row in method_data.iterrows():
        doc.add_paragraph(f"Node: {row['Node']}, Avg Improvement Time: {row['Improvement Time (Difference)']} ms")

# Add perturbative average improvement time to the document
doc.add_heading("Perturbative Methods", level=2)
for method in perturbative_methods:
    doc.add_heading(f"Method: {method}", level=3)
    method_data = perturbative_avg_time[perturbative_avg_time['Improvement Heuristics'] == method]
    for _, row in method_data.iterrows():
        doc.add_paragraph(f"Node: {row['Node']}, Avg Improvement Time: {row['Improvement Time (Difference)']} ms")

# Plot graphs
# Constructive methods: Average Improvement Time vs Node
plt.figure(figsize=(10, 6))
for method in constructive_methods:
    method_data = constructive_avg_time[constructive_avg_time['Construction Heuristics'] == method]
    plt.plot(method_data['Node'], method_data['Improvement Time (Difference)'], marker='o', label=method)
plt.title("Constructive Methods: Avg Improvement Time vs Node")
plt.xlabel("Node Number")
plt.ylabel("Avg Improvement Time (ms)")
plt.grid()
plt.legend()
constructive_time_plot_path = "constructive_avg_improvement_time.png"
plt.savefig(constructive_time_plot_path)
plt.close()
doc.add_picture(constructive_time_plot_path, width=Inches(5))

# Perturbative methods: Average Improvement Time vs Node
plt.figure(figsize=(10, 6))
for method in perturbative_methods:
    method_data = perturbative_avg_time[perturbative_avg_time['Improvement Heuristics'] == method]
    plt.plot(method_data['Node'], method_data['Improvement Time (Difference)'], marker='o', label=method)
plt.title("Perturbative Methods: Avg Improvement Time vs Node")
plt.xlabel("Node Number")
plt.ylabel("Avg Improvement Time (ms)")
plt.grid()
plt.legend()
perturbative_time_plot_path = "perturbative_avg_improvement_time.png"
plt.savefig(perturbative_time_plot_path)
plt.close()
doc.add_picture(perturbative_time_plot_path, width=Inches(5))

# Save the document
doc.save("TSP_Improvement_Time_Analysis.docx")

print("Document created successfully: TSP_Improvement_Time_Analysis.docx")



